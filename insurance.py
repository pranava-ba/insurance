"""
Insurance Churn - Confirmatory Data Analysis (CDA) Streamlit App
=================================================================
Requires: streamlit, pandas, numpy, scipy, pingouin, matplotlib,
          seaborn, python-docx, openpyxl
"""

import io
import re
import textwrap
import warnings


import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy import stats
from scipy.stats import (
    chi2_contingency, mannwhitneyu, f_oneway,
    pearsonr, spearmanr, shapiro, levene
)
warnings.filterwarnings("ignore")

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Insurance Churn – CDA",
    page_icon="📊",
    layout="wide",
)

st.markdown("""
<style>
.main-header {font-size:2rem; font-weight:700; color:#1a3a5c; margin-bottom:0.2rem;}
.section-header {font-size:1.15rem; font-weight:600; color:#1a3a5c;
                 border-bottom:2px solid #1a3a5c; padding-bottom:4px; margin-top:1.2rem;}
.step-box {background:#f0f4fa; border-left:4px solid #1a73e8;
           padding:0.6rem 1rem; border-radius:4px; margin-bottom:0.8rem;}
.result-card {background:#fff; border:1px solid #d0d7e3; border-radius:8px;
              padding:1rem 1.4rem; margin-bottom:1.2rem; box-shadow:0 1px 4px rgba(0,0,0,.06);}
.reject {color:#c0392b; font-weight:700;}
.fail-reject {color:#27ae60; font-weight:700;}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

ALPHA = 0.05

def load_data(uploaded_file):
    try:
        df = pd.read_excel(uploaded_file, engine="openpyxl")
        return df, None
    except Exception as e:
        return None, str(e)


def infer_variable_types(df):
    """Classify columns into categorical / continuous."""
    cats, conts = [], []
    for col in df.columns:
        if df[col].dtype == object or df[col].nunique() <= 10:
            cats.append(col)
        else:
            conts.append(col)
    return cats, conts


def run_analysis(question: str, df: pd.DataFrame):
    """
    Parse the research question and select the appropriate statistical test.
    Returns a dict with all 5 CDA steps + optional figure bytes.
    """
    cats, conts = infer_variable_types(df)
    col_lower = {c.lower(): c for c in df.columns}

    # ── detect columns mentioned in question ──
    mentioned = []
    for col in df.columns:
        if col.lower() in question.lower():
            mentioned.append(col)
    # fallback: guess from keywords
    if not mentioned:
        for col in df.columns:
            words = re.findall(r'\b\w+\b', col.lower())
            if any(w in question.lower() for w in words if len(w) > 3):
                mentioned.append(col)

    # deduplicate preserving order
    seen = set(); mentioned = [x for x in mentioned if not (x in seen or seen.add(x))]

    result = {
        "question": question,
        "columns_used": mentioned,
        "steps": {},
        "fig_bytes": None,
        "test_name": "",
        "error": None,
    }

    try:
        # ── decide test ──
        if len(mentioned) < 2:
            result["error"] = (
                "Could not reliably identify two variables from the question. "
                "Please include exact column names as they appear in the dataset."
            )
            return result

        c1, c2 = mentioned[0], mentioned[1]
        t1 = "cat" if c1 in cats else "cont"
        t2 = "cat" if c2 in cats else "cont"

        data = df[[c1, c2]].dropna()

        # ── CHI-SQUARE: cat vs cat ──
        if t1 == "cat" and t2 == "cat":
            result["test_name"] = "Chi-Square Test of Independence"
            ct = pd.crosstab(data[c1], data[c2])
            chi2, p, dof, expected = chi2_contingency(ct)
            critical = stats.chi2.ppf(1 - ALPHA, dof)

            result["steps"] = {
                "H0": f"There is NO significant association between **{c1}** and **{c2}**.",
                "H1": f"There IS a significant association between **{c1}** and **{c2}**.",
                "test_stat": (
                    f"**Test:** Chi-Square (χ²) Test of Independence  \n"
                    f"**χ² statistic** = {chi2:.4f}  \n"
                    f"**Degrees of freedom** = {dof}  \n"
                    f"**Contingency Table:**\n{ct.to_markdown()}"
                ),
                "sig_level": (
                    f"**Level of significance (α)** = {ALPHA}  \n"
                    f"**p-value** = {p:.6f}  \n"
                    f"**Critical value** (χ²[{dof}, α={ALPHA}]) = {critical:.4f}"
                ),
                "decision": _decision(p, ALPHA, chi2, critical, "χ²"),
            }
            # plot
            fig, ax = plt.subplots(figsize=(7, 4))
            ct_pct = ct.div(ct.sum(axis=1), axis=0) * 100
            ct_pct.plot(kind="bar", ax=ax, colormap="Set2", edgecolor="white")
            ax.set_title(f"{c1} vs {c2} (row %)")
            ax.set_ylabel("Percentage (%)")
            ax.legend(title=c2, bbox_to_anchor=(1, 1))
            plt.tight_layout()
            result["fig_bytes"] = _fig_to_bytes(fig)

        # ── INDEPENDENT T-TEST / MANN-WHITNEY: cat(2) vs cont ──
        elif (t1 == "cat" and t2 == "cont") or (t1 == "cont" and t2 == "cat"):
            cat_col = c1 if t1 == "cat" else c2
            cont_col = c2 if t1 == "cat" else c1
            groups = data[cat_col].unique()

            if len(groups) == 2:
                g1 = data.loc[data[cat_col] == groups[0], cont_col]
                g2 = data.loc[data[cat_col] == groups[1], cont_col]
                # normality check
                _, p_norm1 = shapiro(g1.sample(min(len(g1), 200), random_state=42))
                _, p_norm2 = shapiro(g2.sample(min(len(g2), 200), random_state=42))
                normal = p_norm1 > 0.05 and p_norm2 > 0.05

                if normal:
                    result["test_name"] = "Independent Samples t-Test"
                    _, p_lev = levene(g1, g2)
                    equal_var = p_lev > 0.05
                    t_stat, p = stats.ttest_ind(g1, g2, equal_var=equal_var)
                    dof = len(g1) + len(g2) - 2
                    critical = stats.t.ppf(1 - ALPHA / 2, dof)
                    result["steps"] = {
                        "H0": f"There is NO significant difference in mean **{cont_col}** between the groups of **{cat_col}**.",
                        "H1": f"There IS a significant difference in mean **{cont_col}** between the groups of **{cat_col}**.",
                        "test_stat": (
                            f"**Test:** Independent Samples t-Test ({'Equal' if equal_var else 'Unequal'} variances – Welch's)  \n"
                            f"**Group '{groups[0]}'**: n={len(g1)}, mean={g1.mean():.2f}, sd={g1.std():.2f}  \n"
                            f"**Group '{groups[1]}'**: n={len(g2)}, mean={g2.mean():.2f}, sd={g2.std():.2f}  \n"
                            f"**t-statistic** = {t_stat:.4f}  \n"
                            f"**Degrees of freedom** = {dof}"
                        ),
                        "sig_level": (
                            f"**Level of significance (α)** = {ALPHA}  \n"
                            f"**p-value** = {p:.6f}  \n"
                            f"**Critical value** t[{dof}, α/2={ALPHA/2}] = ±{critical:.4f}"
                        ),
                        "decision": _decision(p, ALPHA, abs(t_stat), critical, "t"),
                    }
                else:
                    result["test_name"] = "Mann-Whitney U Test (non-parametric)"
                    u_stat, p = mannwhitneyu(g1, g2, alternative="two-sided")
                    result["steps"] = {
                        "H0": f"There is NO significant difference in **{cont_col}** between the groups of **{cat_col}**.",
                        "H1": f"There IS a significant difference in **{cont_col}** between the groups of **{cat_col}**.",
                        "test_stat": (
                            f"**Test:** Mann-Whitney U Test (normality assumption violated)  \n"
                            f"**Group '{groups[0]}'**: n={len(g1)}, median={g1.median():.2f}  \n"
                            f"**Group '{groups[1]}'**: n={len(g2)}, median={g2.median():.2f}  \n"
                            f"**U-statistic** = {u_stat:.2f}"
                        ),
                        "sig_level": (
                            f"**Level of significance (α)** = {ALPHA}  \n"
                            f"**p-value** = {p:.6f}"
                        ),
                        "decision": _decision_p(p, ALPHA),
                    }

            else:
                # >2 groups: ANOVA or Kruskal-Wallis
                group_data = [data.loc[data[cat_col] == g, cont_col] for g in groups]
                _, p_norm = shapiro(data[cont_col].sample(min(len(data), 200), random_state=42))
                if p_norm > 0.05:
                    result["test_name"] = "One-Way ANOVA"
                    f_stat, p = f_oneway(*group_data)
                    dof_b = len(groups) - 1
                    dof_w = len(data) - len(groups)
                    critical = stats.f.ppf(1 - ALPHA, dof_b, dof_w)
                    means_str = "  \n".join([f"  - {g}: mean={d.mean():.2f}, n={len(d)}" for g, d in zip(groups, group_data)])
                    result["steps"] = {
                        "H0": f"There is NO significant difference in mean **{cont_col}** across all groups of **{cat_col}**.",
                        "H1": f"At least one group mean of **{cont_col}** differs significantly across **{cat_col}** groups.",
                        "test_stat": (
                            f"**Test:** One-Way ANOVA  \n"
                            f"**Group means:**  \n{means_str}  \n"
                            f"**F-statistic** = {f_stat:.4f}  \n"
                            f"**df between** = {dof_b}, **df within** = {dof_w}"
                        ),
                        "sig_level": (
                            f"**Level of significance (α)** = {ALPHA}  \n"
                            f"**p-value** = {p:.6f}  \n"
                            f"**Critical value** F[{dof_b},{dof_w}] = {critical:.4f}"
                        ),
                        "decision": _decision(p, ALPHA, f_stat, critical, "F"),
                    }
                else:
                    result["test_name"] = "Kruskal-Wallis Test (non-parametric)"
                    h_stat, p = stats.kruskal(*group_data)
                    result["steps"] = {
                        "H0": f"There is NO significant difference in **{cont_col}** across groups of **{cat_col}**.",
                        "H1": f"At least one group of **{cat_col}** has a significantly different distribution of **{cont_col}**.",
                        "test_stat": (
                            f"**Test:** Kruskal-Wallis H Test  \n"
                            f"**H-statistic** = {h_stat:.4f}  \n"
                            f"**Groups tested:** {list(groups)}"
                        ),
                        "sig_level": (
                            f"**Level of significance (α)** = {ALPHA}  \n"
                            f"**p-value** = {p:.6f}"
                        ),
                        "decision": _decision_p(p, ALPHA),
                    }

            # box plot
            fig, ax = plt.subplots(figsize=(7, 4))
            cats_order = sorted(data[cat_col].unique())
            sns.boxplot(data=data, x=cat_col, y=cont_col, order=cats_order,
                        palette="Set2", ax=ax)
            ax.set_title(f"{cont_col} by {cat_col}")
            plt.xticks(rotation=30, ha="right")
            plt.tight_layout()
            result["fig_bytes"] = _fig_to_bytes(fig)

        # ── CORRELATION: cont vs cont ──
        elif t1 == "cont" and t2 == "cont":
            _, p_n1 = shapiro(data[c1].sample(min(len(data), 200), random_state=42))
            _, p_n2 = shapiro(data[c2].sample(min(len(data), 200), random_state=42))
            normal = p_n1 > 0.05 and p_n2 > 0.05

            if normal:
                result["test_name"] = "Pearson Correlation"
                r, p = pearsonr(data[c1], data[c2])
                n = len(data)
                t_stat = r * np.sqrt(n - 2) / np.sqrt(1 - r**2)
                dof = n - 2
                critical = stats.t.ppf(1 - ALPHA / 2, dof)
                result["steps"] = {
                    "H0": f"There is NO significant linear correlation between **{c1}** and **{c2}** (ρ = 0).",
                    "H1": f"There IS a significant linear correlation between **{c1}** and **{c2}** (ρ ≠ 0).",
                    "test_stat": (
                        f"**Test:** Pearson Product-Moment Correlation  \n"
                        f"**r** = {r:.4f}  \n"
                        f"**t-statistic** = {t_stat:.4f}  \n"
                        f"**Degrees of freedom** = {dof}  \n"
                        f"**n** = {n}"
                    ),
                    "sig_level": (
                        f"**Level of significance (α)** = {ALPHA}  \n"
                        f"**p-value** = {p:.6f}  \n"
                        f"**Critical value** t[{dof}] = ±{critical:.4f}"
                    ),
                    "decision": _decision(p, ALPHA, abs(t_stat), critical, "t"),
                }
            else:
                result["test_name"] = "Spearman Rank Correlation"
                rho, p = spearmanr(data[c1], data[c2])
                n = len(data)
                result["steps"] = {
                    "H0": f"There is NO significant monotonic relationship between **{c1}** and **{c2}** (ρₛ = 0).",
                    "H1": f"There IS a significant monotonic relationship between **{c1}** and **{c2}** (ρₛ ≠ 0).",
                    "test_stat": (
                        f"**Test:** Spearman Rank Correlation (non-parametric)  \n"
                        f"**ρₛ (Spearman rho)** = {rho:.4f}  \n"
                        f"**n** = {n}"
                    ),
                    "sig_level": (
                        f"**Level of significance (α)** = {ALPHA}  \n"
                        f"**p-value** = {p:.6f}"
                    ),
                    "decision": _decision_p(p, ALPHA),
                }

            # scatter
            fig, ax = plt.subplots(figsize=(6, 4))
            ax.scatter(data[c1], data[c2], alpha=0.3, s=15, color="#1a73e8")
            m, b = np.polyfit(data[c1], data[c2], 1)
            ax.plot(data[c1], m * data[c1] + b, color="#c0392b", lw=1.5)
            ax.set_xlabel(c1); ax.set_ylabel(c2)
            ax.set_title(f"{c1} vs {c2}")
            plt.tight_layout()
            result["fig_bytes"] = _fig_to_bytes(fig)

    except Exception as exc:
        result["error"] = str(exc)

    return result


def _decision(p, alpha, stat, critical, stat_name):
    if p < alpha:
        return (
            f"Since **p-value ({p:.4f}) < α ({alpha})** and "
            f"|{stat_name}| = {stat:.4f} > critical value {critical:.4f}, "
            f"we **reject the null hypothesis**. "
            f"There is sufficient statistical evidence to support the alternate hypothesis."
        )
    else:
        return (
            f"Since **p-value ({p:.4f}) ≥ α ({alpha})** and "
            f"|{stat_name}| = {stat:.4f} ≤ critical value {critical:.4f}, "
            f"we **fail to reject the null hypothesis**. "
            f"There is insufficient statistical evidence to support the alternate hypothesis."
        )


def _decision_p(p, alpha):
    if p < alpha:
        return (
            f"Since **p-value ({p:.4f}) < α ({alpha})**, "
            f"we **reject the null hypothesis**. "
            f"There is sufficient statistical evidence to support the alternate hypothesis."
        )
    else:
        return (
            f"Since **p-value ({p:.4f}) ≥ α ({alpha})**, "
            f"we **fail to reject the null hypothesis**. "
            f"There is insufficient statistical evidence to support the alternate hypothesis."
        )


def _fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# WORD EXPORT
# ─────────────────────────────────────────────

def build_word_report(results, df_info):
    doc = Document()

    # Title
    title = doc.add_heading("Insurance Churn – Confirmatory Data Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph(
        f"Dataset: {df_info['rows']} observations × {df_info['cols']} variables  |  "
        f"Level of Significance (α) = 0.05"
    ).runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    for i, res in enumerate(results, 1):
        doc.add_heading(f"Question {i}: {res['question']}", level=1)

        if res.get("error"):
            doc.add_paragraph(f"⚠ Error: {res['error']}")
            continue

        doc.add_paragraph(f"Variables identified: {', '.join(res['columns_used'])}")
        doc.add_paragraph(f"Statistical Test Applied: {res['test_name']}")
        doc.add_paragraph("")

        steps_map = {
            "H0": "Step 1 – Null Hypothesis (H₀)",
            "H1": "Step 2 – Alternate Hypothesis (H₁)",
            "test_stat": "Step 3 – Test Statistic",
            "sig_level": "Step 4 – Level of Significance & p-value",
            "decision": "Step 5 – Decision",
        }

        for key, label in steps_map.items():
            h = doc.add_heading(label, level=2)
            h.runs[0].font.size = Pt(12)
            # strip markdown bold markers for Word
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', res["steps"].get(key, ""))
            doc.add_paragraph(text)

        if res.get("fig_bytes"):
            doc.add_paragraph("")
            img_stream = io.BytesIO(res["fig_bytes"])
            doc.add_picture(img_stream, width=Inches(5.5))

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# APP LAYOUT
# ─────────────────────────────────────────────

st.markdown('<div class="main-header">📊 Insurance Churn – Confirmatory Data Analysis</div>', unsafe_allow_html=True)
st.markdown("Conduct rigorous, hypothesis-driven statistical tests on the insurance churn dataset.")

# ── SESSION STATE ──
if "df" not in st.session_state:
    st.session_state.df = None
if "results" not in st.session_state:
    st.session_state.results = []

# ═══════════════════════════════════════════
# STEP 1 – LOAD DATA
# ═══════════════════════════════════════════
st.markdown("---")
st.markdown('<div class="section-header">Step 1 · Load Dataset</div>', unsafe_allow_html=True)

upload = st.file_uploader("Upload the Insurance Churn Excel file (.xlsx)", type=["xlsx"])

if upload:
    df, err = load_data(upload)
    if err:
        st.error(f"Could not load file: {err}")
    else:
        st.session_state.df = df
        st.success(f"✅ Dataset loaded — **{df.shape[0]:,} rows × {df.shape[1]} columns**")

        col1, col2 = st.columns([2, 1])
        with col1:
            st.dataframe(df.head(10), use_container_width=True)
        with col2:
            cats, conts = infer_variable_types(df)
            st.markdown("**Variable Summary**")
            st.markdown(f"- 🗂 Categorical variables ({len(cats)}): {', '.join(cats)}")
            st.markdown(f"- 📈 Continuous variables ({len(conts)}): {', '.join(conts)}")
            st.markdown(f"- 🎯 Target variable: **Churn** (0 = stayed, 1 = churned)")
            st.markdown(f"- Churn rate: **{df['Churn'].mean()*100:.1f}%**")

# ═══════════════════════════════════════════
# STEP 2 – CONFIGURE QUESTIONS
# ═══════════════════════════════════════════
if st.session_state.df is not None:
    df = st.session_state.df
    st.markdown("---")
    st.markdown('<div class="section-header">Step 2 · Define Research Questions</div>', unsafe_allow_html=True)

    st.markdown(
        '<div class="step-box">Enter each research question using the <b>exact column names</b> '
        'from the dataset. Example: <i>"Is there a significant association between Gender and Churn?"</i></div>',
        unsafe_allow_html=True,
    )

    n_qs = st.selectbox(
        "How many hypotheses do you want to test?",
        options=list(range(2, 7)),
        index=0,
        help="Choose between 2 and 6 questions."
    )

    st.markdown("**Column names for reference:**")
    st.code(", ".join(df.columns.tolist()), language=None)

    questions = []
    for i in range(n_qs):
        q = st.text_area(
            f"Question {i+1}",
            placeholder=f"e.g. Is there a significant difference in Credit Score between Churn groups?",
            height=68,
            key=f"q_{i}",
        )
        questions.append(q.strip())

# ═══════════════════════════════════════════
# STEP 3 – RUN ANALYSIS
# ═══════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="section-header">Step 3 · Run Confirmatory Data Analysis</div>', unsafe_allow_html=True)

    run_btn = st.button("🔬 Run Analysis", type="primary", use_container_width=True)

    if run_btn:
        filled = [q for q in questions if q]
        if len(filled) < 2:
            st.warning("Please enter at least 2 questions before running the analysis.")
        else:
            st.session_state.results = []
            with st.spinner("Running statistical tests…"):
                for q in filled:
                    res = run_analysis(q, df)
                    st.session_state.results.append(res)

    # ── DISPLAY RESULTS ──
    if st.session_state.results:
        for i, res in enumerate(st.session_state.results, 1):
            st.markdown(f"---")
            st.markdown(f"### 🔍 Question {i}: _{res['question']}_")

            if res.get("error"):
                st.error(f"⚠ {res['error']}")
                st.info(
                    "Tip: Make sure your question contains **exact column names** "
                    "as they appear in the dataset header."
                )
                continue

            st.markdown(
                f'<div class="result-card">'
                f'<b>Variables detected:</b> {", ".join(res["columns_used"])}  &nbsp;|&nbsp;  '
                f'<b>Test applied:</b> {res["test_name"]}'
                f'</div>',
                unsafe_allow_html=True,
            )

            step_labels = {
                "H0": ("1", "Null Hypothesis (H₀)"),
                "H1": ("2", "Alternate Hypothesis (H₁)"),
                "test_stat": ("3", "Test Statistic"),
                "sig_level": ("4", "Level of Significance & p-value"),
                "decision": ("5", "Decision"),
            }

            for key, (num, label) in step_labels.items():
                text = res["steps"].get(key, "")
                st.markdown(f"**Step {num} – {label}**")
                # colour decision box
                if key == "decision":
                    if "reject the null" in text and "fail to" not in text:
                        st.markdown(
                            f'<div class="result-card"><span class="reject">🔴 {text}</span></div>',
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(
                            f'<div class="result-card"><span class="fail-reject">🟢 {text}</span></div>',
                            unsafe_allow_html=True,
                        )
                else:
                    st.markdown(text)
                st.markdown("")

            if res.get("fig_bytes"):
                st.image(res["fig_bytes"], use_column_width=False, width=650)

        # ── DOWNLOAD BUTTON ──
        st.markdown("---")
        st.markdown('<div class="section-header">Download Results</div>', unsafe_allow_html=True)

        df_info = {"rows": df.shape[0], "cols": df.shape[1]}
        word_bytes = build_word_report(st.session_state.results, df_info)

        st.download_button(
            label="📄 Download Full Report as Word Document (.docx)",
            data=word_bytes,
            file_name="Insurance_Churn_CDA_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
